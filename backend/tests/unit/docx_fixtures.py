"""Building Word documents that carry comments, for tests.

python-docx models no part of a comment, so these three parts are written by
hand. That is not a shortcut — it is the only way to get the shape a real
reviewer's file has: a thread with a reply, marked resolved, anchored over a
phrase in the body rather than over a whole paragraph.

Not named ``test_*`` so pytest does not try to collect it.
"""

from __future__ import annotations

import io
import zipfile

import docx

COMMENTS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
  <w:comment w:id="0" w:author="Priya Raman" w:initials="PR" w:date="2026-08-01T10:00:00Z">
    <w:p w14:paraId="AAAA0001"><w:r><w:t>Is this still the enterprise price?</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="1" w:author="Sam Okafor" w:initials="SO" w:date="2026-08-01T11:00:00Z">
    <w:p w14:paraId="AAAA0002"><w:r><w:t>Checking with finance.</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="2" w:author="Priya Raman" w:initials="PR">
    <w:p w14:paraId="AAAA0003"><w:r><w:t>Tighten this sentence.</w:t></w:r></w:p>
  </w:comment>
</w:comments>
"""

# Thread 0 is resolved and 1 replies to it; comment 2 is open and unthreaded.
COMMENTS_EX_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">
  <w15:commentEx w15:paraId="AAAA0001" w15:done="1"/>
  <w15:commentEx w15:paraId="AAAA0002" w15:paraIdParent="AAAA0001" w15:done="1"/>
  <w15:commentEx w15:paraId="AAAA0003" w15:done="0"/>
</w15:commentsEx>
"""

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def anchor(document_xml: bytes, phrase: str, comment_id: str) -> bytes:
    """Anchor ``comment_id`` over exactly ``phrase``.

    Splits the run at the phrase boundaries first, which is what Word does and
    what makes this worth testing: python-docx writes a whole sentence as one
    run, so wrapping the run wholesale would produce an anchor covering the
    paragraph and the test would pass without the range walk doing anything.
    """
    import copy as _copy

    from docx.oxml import parse_xml
    from lxml import etree

    w = f"{{{_W}}}"
    root = parse_xml(document_xml)
    for text_node in root.iter(f"{w}t"):
        whole = text_node.text or ""
        if phrase not in whole:
            continue

        run = text_node.getparent()
        paragraph = run.getparent()
        index = list(paragraph).index(run)
        head, _, tail = whole.partition(phrase)

        def part(text: str, template: etree._Element = run) -> etree._Element:
            clone = _copy.deepcopy(template)
            node = clone.find(f"{w}t")
            node.text = text
            node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            return clone

        pieces: list[etree._Element] = []
        if head:
            pieces.append(part(head))

        start = etree.Element(f"{w}commentRangeStart")
        start.set(f"{w}id", comment_id)
        pieces.append(start)
        pieces.append(part(phrase))
        end = etree.Element(f"{w}commentRangeEnd")
        end.set(f"{w}id", comment_id)
        pieces.append(end)

        reference_run = etree.Element(f"{w}r")
        reference = etree.SubElement(reference_run, f"{w}commentReference")
        reference.set(f"{w}id", comment_id)
        pieces.append(reference_run)

        if tail:
            pieces.append(part(tail))

        paragraph.remove(run)
        for offset, piece in enumerate(pieces):
            paragraph.insert(index + offset, piece)
        break
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def commented_docx() -> bytes:
    """A document carrying a resolved thread with a reply, plus an open remark."""
    document = docx.Document()
    document.add_heading("Pricing", 1)
    document.add_paragraph("The platform costs $50k per year for the enterprise tier.")
    document.add_paragraph("Support is best effort during business hours.")

    plain = io.BytesIO()
    document.save(plain)

    source = zipfile.ZipFile(io.BytesIO(plain.getvalue()))
    out = io.BytesIO()
    with source, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            payload = source.read(item.filename)
            if item.filename == "word/document.xml":
                payload = anchor(payload, "$50k", "0")
                payload = anchor(payload, "best effort", "2")
            elif item.filename == "[Content_Types].xml":
                # Declared so the fixture is a document Word would open, not
                # only one this reader happens to accept.
                payload = payload.replace(
                    b"</Types>",
                    b'<Override PartName="/word/comments.xml" ContentType="application/'
                    b'vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                    b'<Override PartName="/word/commentsExtended.xml" ContentType="'
                    b"application/vnd.ms-word.commentsExtended+xml\"/></Types>",
                )
            target.writestr(item, payload)
        target.writestr("word/comments.xml", COMMENTS_XML)
        target.writestr("word/commentsExtended.xml", COMMENTS_EX_XML)
    return out.getvalue()



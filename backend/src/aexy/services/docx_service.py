"""Word documents, read and written at full fidelity.

Three jobs, one module, because they share a vocabulary and must not disagree
about it:

  * ``extract_structured`` — a ``.docx`` becomes Markdown an LLM can actually
    reason about. This replaces reading ``doc.paragraphs``, which is what the
    file pipeline did and which silently drops **everything inside a table**:
    ``paragraphs`` walks only the top level of the body, so a requirements
    matrix or a pricing grid extracted to nothing at all. Heading levels, list
    nesting, and headers/footers were lost the same way. A summary built from
    that input is confidently wrong rather than obviously empty, which is the
    worse failure.

  * ``render_docx`` — Markdown or a TipTap document becomes ``.docx`` bytes.
    Markdown input goes through the existing ``markdown_to_tiptap``, so the
    codebase keeps one Markdown parser and this module only ever renders the
    editor's own node vocabulary. The single exception is pipe tables, which
    that parser deliberately does not model: see ``_markdown_to_tree``.

  * ``apply_ops`` — a structured edit list is applied to existing bytes. This
    is the *restricted* automation backend: see ``DocxAutomationPort``.

What this module deliberately does not do
-----------------------------------------
It does not write tracked changes. A redline in OOXML is ``w:ins``/``w:del``
wrapping runs, with deleted text relocated into ``w:delText`` and consistent
author/timestamp bookkeeping; python-docx has no API for any of it, and
hand-rolling that XML forfeits the lossless round-trip that the browser
editor's engine guarantees. ``apply_ops`` therefore refuses
``track_changes=True`` rather than producing a document that claims to have a
redline and does not. Unattended redlines are what would justify the Node
sidecar described in the DOCX plan; ``DocxAutomationPort`` is the seam it
would arrive behind.

Known limits, stated rather than discovered
-------------------------------------------
* Vertically or horizontally merged table cells repeat their text across the
  span, because that is what python-docx's ``row.cells`` reports.
* Tables nested inside table cells are extracted as their cell's text, not as
  their own table.
* Footnotes and endnotes are not extracted; python-docx does not expose them.
* Extraction escapes only the characters that would break the structure it
  emits (pipes in table cells). Prose keeps literal ``*`` and ``_`` because
  the consumer is a language model, and over-escaped text reads worse than
  slightly ambiguous text.
"""

from __future__ import annotations

import copy
import io
import logging
import re
import zipfile
from dataclasses import dataclass, field
from dataclasses import replace as _dc_replace
from typing import Any, NoReturn, Protocol

from aexy.services.markdown_to_tiptap import MarkdownError, markdown_to_tiptap

logger = logging.getLogger(__name__)

try:  # pragma: no cover - exercised by the availability guard, not by tests
    from docx import Document as _new_document
    from docx.opc.constants import RELATIONSHIP_TYPE as _REL
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.table import Table as _Table
    from docx.text.paragraph import Paragraph as _Paragraph

    PYTHON_DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    PYTHON_DOCX_AVAILABLE = False


class DocxReadError(ValueError):
    """The bytes handed in are not a readable Word document."""


class DocxRenderError(ValueError):
    """The content handed in could not become a Word document."""


class DocxOpUnsupported(Exception):
    """An edit this backend cannot perform faithfully.

    Raised rather than approximated. A caller that gets bytes back must be
    able to trust that every op in the list was applied as written; a partial
    or best-effort edit to a contract is worse than a refusal.

    Every raise is logged with the stable marker ``docx.op_unsupported`` so
    the rate is greppable. That rate is the signal for whether the restricted
    op set has started to constrain what agents are allowed to propose — which
    is the trigger for moving automation to the sidecar backend.
    """


# ─── Extraction ───────────────────────────────────────────────────────────


@dataclass(frozen=True)
class DocxHeading:
    """One entry in a document's outline."""

    level: int
    text: str


@dataclass(frozen=True)
class DocxTable:
    """A table as a rectangle of cell text, first row first."""

    rows: list[list[str]] = field(default_factory=list)

    @property
    def header(self) -> list[str]:
        return self.rows[0] if self.rows else []


@dataclass(frozen=True)
class DocxParagraph:
    """One paragraph as the edit ops address it.

    ``text`` is the run-joined string, which is exactly what a ``replace_text``
    op's ``find`` is compared against. That is the whole reason this exists
    beside ``markdown``: a model prompted with Markdown proposes
    ``find: "**Tier**"`` for a bolded cell, which matches nothing, and the
    reviewer opens an empty redline with nothing to explain it. Prompt with
    these instead and every ``find`` the model writes is a string the document
    actually contains.

    ``index`` counts non-empty paragraphs in document order, so it is a label
    for a prompt ("paragraph 12") rather than an address — no op takes one.
    """

    index: int
    text: str
    style: str
    heading_level: int | None = None
    in_table: bool = False


@dataclass(frozen=True)
class DocxExtract:
    """Everything a reader — human or model — needs from a Word document.

    ``markdown`` is the form to hand an LLM and to persist as a document's
    searchable text. ``plain_text`` is the same content with the structural
    markup removed, for word counts and keyword search. ``outline`` and
    ``tables`` are addressable structure: intake keys proposed work items to
    the heading they came from, so a candidate task can point back at the
    section that produced it. ``paragraphs`` is the form to prompt an editing
    model with — see ``DocxParagraph`` for why Markdown will not do.
    """

    markdown: str
    plain_text: str
    outline: list[DocxHeading] = field(default_factory=list)
    tables: list[DocxTable] = field(default_factory=list)
    paragraphs: list[DocxParagraph] = field(default_factory=list)
    word_count: int = 0


_HEADING_STYLE = re.compile(r"^heading\s*([1-9])$")
_HEADING_STYLE_ID = re.compile(r"^heading([1-9])$")


def _require_python_docx() -> None:
    if not PYTHON_DOCX_AVAILABLE:
        raise DocxReadError(
            "Word document support requires python-docx. "
            "Install with: pip install python-docx"
        )


def _iter_block_items(parent: Any, container: Any) -> Any:
    """Yield paragraphs and tables in the order they appear in the document.

    The whole point of this module. ``document.paragraphs`` returns only the
    body's top-level ``w:p`` children, so table content is invisible to it and
    a paragraph that follows a table cannot be distinguished from one that
    precedes it. Walking the XML children in order is the only way to get
    document order, and document order is what makes the extracted Markdown
    mean the same thing as the document.
    """
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield _Paragraph(child, container)
        elif child.tag == qn("w:tbl"):
            yield _Table(child, container)


def _style_name(paragraph: Any) -> str:
    style = getattr(paragraph, "style", None)
    return (getattr(style, "name", "") or "").strip()


def _outline_level(paragraph: Any) -> int | None:
    """Heading level from ``w:outlineLvl``, for documents with renamed styles.

    A document authored in a non-English Word has styles named "Titre 1" or
    "शीर्षक 1", and a document built from a corporate template often has
    "AcmeHeading". Both still carry ``outlineLvl``, so this is the fallback
    that keeps a translated document from extracting as flat prose.
    """
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    element = pPr.find(qn("w:outlineLvl"))
    if element is None:
        return None
    raw = element.get(qn("w:val"))
    try:
        level = int(raw)
    except (TypeError, ValueError):
        return None
    return level + 1 if 0 <= level <= 8 else None


def _heading_level(paragraph: Any) -> int | None:
    """The heading level of a paragraph, or None if it is not a heading."""
    name = _style_name(paragraph).lower()

    match = _HEADING_STYLE.match(name)
    if match:
        return int(match.group(1))

    style_id = (getattr(getattr(paragraph, "style", None), "style_id", "") or "").lower()
    match = _HEADING_STYLE_ID.match(style_id)
    if match:
        return int(match.group(1))

    # "Title" and "Subtitle" are structural in every template that ships them,
    # and a document whose only structure is a Title extracts as one flat
    # blob without this.
    if name == "title":
        return 1
    if name == "subtitle":
        return 2

    return _outline_level(paragraph)


def _list_marker(
    paragraph: Any, numbering: dict[tuple[int, int], str]
) -> tuple[int, bool, int | None] | None:
    """``(indent_level, is_ordered, num_id)`` for a list paragraph, else None."""
    pPr = paragraph._p.pPr
    numPr = pPr.numPr if pPr is not None else None

    if numPr is None:
        # A paragraph styled as a list but carrying no numbering reference
        # still reads as a list to a human, so it should to a model too.
        name = _style_name(paragraph).lower()
        if name.startswith("list number"):
            return 0, True, None
        if name.startswith(("list bullet", "list paragraph")):
            return 0, False, None
        return None

    ilvl = 0
    if numPr.ilvl is not None and numPr.ilvl.val is not None:
        ilvl = int(numPr.ilvl.val)
    num_id = None
    if numPr.numId is not None and numPr.numId.val is not None:
        num_id = int(numPr.numId.val)

    fmt = numbering.get((num_id, ilvl)) if num_id is not None else None
    if fmt is None:
        ordered = _style_name(paragraph).lower().startswith("list number")
    else:
        ordered = fmt not in {"bullet", "none"}

    return ilvl, ordered, num_id


def _numbering_formats(document: Any) -> dict[tuple[int, int], str]:
    """Map ``(numId, ilvl)`` to its ``w:numFmt``, from ``numbering.xml``.

    Without this, an ordered list and a bulleted list are indistinguishable:
    both are just paragraphs carrying a ``numPr``. Numbered steps extracted as
    bullets lose the fact that they are ordered, which for a procedure is the
    whole meaning.

    Best-effort by design — a document with no numbering part, or one whose
    numbering is defined in a way this does not model, falls back to the
    style-name heuristic rather than failing the extraction.
    """
    try:
        numbering = document.part.numbering_part.element
    except (AttributeError, KeyError, ValueError):
        return {}

    abstract: dict[int, dict[int, str]] = {}
    for node in numbering.findall(qn("w:abstractNum")):
        raw_id = node.get(qn("w:abstractNumId"))
        try:
            abstract_id = int(raw_id)
        except (TypeError, ValueError):
            continue
        levels: dict[int, str] = {}
        for lvl in node.findall(qn("w:lvl")):
            try:
                ilvl = int(lvl.get(qn("w:ilvl")))
            except (TypeError, ValueError):
                continue
            fmt_node = lvl.find(qn("w:numFmt"))
            if fmt_node is not None:
                levels[ilvl] = (fmt_node.get(qn("w:val")) or "").lower()
        abstract[abstract_id] = levels

    formats: dict[tuple[int, int], str] = {}
    for node in numbering.findall(qn("w:num")):
        try:
            num_id = int(node.get(qn("w:numId")))
        except (TypeError, ValueError):
            continue
        ref = node.find(qn("w:abstractNumId"))
        if ref is None:
            continue
        try:
            abstract_id = int(ref.get(qn("w:val")))
        except (TypeError, ValueError):
            continue
        for ilvl, fmt in abstract.get(abstract_id, {}).items():
            formats[(num_id, ilvl)] = fmt

    return formats


def _run_markdown(run_element: Any, part: Any) -> str:
    """One ``w:r`` as Markdown, carrying bold and italic."""
    pieces: list[str] = []
    for node in run_element.iter():
        if node.tag == qn("w:t"):
            pieces.append(node.text or "")
        elif node.tag == qn("w:tab"):
            pieces.append("\t")
        elif node.tag in (qn("w:br"), qn("w:cr")):
            pieces.append(" ")
    text = "".join(pieces)
    if not text.strip():
        return text

    rPr = run_element.find(qn("w:rPr"))
    bold = italic = False
    if rPr is not None:
        bold = _toggle_on(rPr.find(qn("w:b")))
        italic = _toggle_on(rPr.find(qn("w:i")))

    # Marks wrap the trimmed text: "**text **" renders literally in most
    # Markdown parsers, so trailing whitespace has to sit outside the marker.
    lead = text[: len(text) - len(text.lstrip())]
    trail = text[len(text.rstrip()) :]
    core = text.strip()
    if bold:
        core = f"**{core}**"
    if italic:
        core = f"*{core}*"
    return f"{lead}{core}{trail}"


def _toggle_on(element: Any) -> bool:
    """Whether an OOXML boolean property is on.

    ``<w:b/>`` means on; ``<w:b w:val="0"/>`` means explicitly off, which is
    how a run inside a bold style turns bold back off.
    """
    if element is None:
        return False
    val = element.get(qn("w:val"))
    return val is None or val not in {"0", "false", "off"}


def _paragraph_markdown(paragraph: Any) -> str:
    """A paragraph's inline content as Markdown, hyperlinks included."""
    part = paragraph.part
    pieces: list[str] = []

    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:r"):
            pieces.append(_run_markdown(child, part))
        elif child.tag == qn("w:hyperlink"):
            label = "".join(
                _run_markdown(run, part) for run in child.findall(qn("w:r"))
            ).strip()
            if not label:
                continue
            rel_id = child.get(qn("r:id"))
            target = None
            if rel_id:
                try:
                    target = part.rels[rel_id].target_ref
                except (KeyError, AttributeError):
                    target = None
            pieces.append(f"[{label}]({target})" if target else label)

    return "".join(pieces).strip()


def _paragraph_text(paragraph: Any) -> str:
    """A paragraph's text with no markup, hyperlink labels included.

    Not ``paragraph.text``: that skips runs inside ``w:hyperlink``, so a line
    that is entirely a link reads as empty.
    """
    pieces: list[str] = []
    for node in paragraph._p.iter():
        if node.tag == qn("w:t"):
            pieces.append(node.text or "")
        elif node.tag == qn("w:tab"):
            pieces.append("\t")
    return "".join(pieces).strip()


_MONOSPACE_FONTS = {
    "consolas",
    "courier",
    "courier new",
    "menlo",
    "monaco",
    "sf mono",
    "cascadia code",
    "cascadia mono",
    "dejavu sans mono",
    "lucida console",
    "source code pro",
    "jetbrains mono",
    "roboto mono",
}


def _is_code_paragraph(paragraph: Any) -> bool:
    """Whether a paragraph is a line of code rather than prose.

    Word has no code block. Every convention for writing one — the "Code"
    style if the template has it, otherwise a monospace font — is a font
    choice, so a font check is the only signal available. Aexy's documents are
    technical, and a config snippet flattened into a paragraph loses the
    indentation that was its meaning.
    """
    name = _style_name(paragraph).lower()
    if "code" in name or name in {"html preformatted", "plain text", "preformatted text"}:
        return True

    runs = [node for node in paragraph._p.iterchildren() if node.tag == qn("w:r")]
    if not runs:
        return False

    saw_font = False
    for run in runs:
        rPr = run.find(qn("w:rPr"))
        fonts = rPr.find(qn("w:rFonts")) if rPr is not None else None
        family = (fonts.get(qn("w:ascii")) or "").lower() if fonts is not None else ""
        if not family:
            return False
        if family not in _MONOSPACE_FONTS:
            return False
        saw_font = True
    return saw_font


def _is_horizontal_rule(paragraph: Any) -> bool:
    """An empty paragraph carrying a bottom border — Word's horizontal rule."""
    if _paragraph_text(paragraph):
        return False
    pPr = paragraph._p.pPr
    if pPr is None:
        return False
    borders = pPr.find(qn("w:pBdr"))
    return borders is not None and borders.find(qn("w:bottom")) is not None


def _is_quote(paragraph: Any) -> bool:
    return "quote" in _style_name(paragraph).lower()


def _cell_text(cell: Any) -> str:
    """A table cell as a single line, safe to place inside a pipe table."""
    parts = [_paragraph_markdown(p) for p in cell.paragraphs]
    text = " ".join(p for p in parts if p)
    return text.replace("|", "\\|").replace("\n", " ").strip()


def _cell_plain(cell: Any) -> str:
    """A table cell's text with no markup."""
    parts = [_paragraph_text(p) for p in cell.paragraphs]
    return " ".join(p for p in parts if p).replace("\n", " ").strip()


def _table_rows(table: Any) -> tuple[list[list[str]], list[list[str]]]:
    """A table's cells twice: as Markdown, and as plain text.

    Two views because they answer different questions. The Markdown view goes
    into the pipe table a model reads, where a bold header should look bold.
    The plain view is what ``DocxExtract.tables`` exposes, because that is
    addressable data — intake keys candidate work items to it and
    ``set_table_cell`` addresses it by coordinate. A caller comparing a cell to
    ``"Task"`` must not have to know it might be ``"**Task**"``.
    """
    try:
        markdown_rows = []
        plain_rows = []
        for row in table.rows:
            cells = list(row.cells)
            markdown_rows.append([_cell_text(cell) for cell in cells])
            plain_rows.append([_cell_plain(cell) for cell in cells])
        return markdown_rows, plain_rows
    except (AttributeError, IndexError, ValueError) as exc:
        logger.warning("Skipping unreadable table: %s", exc)
        return [], []


def _table_markdown(rows: list[list[str]]) -> list[str]:
    """Pipe-table Markdown for a rectangle of cells.

    Every table gets a header separator even when its first row is data.
    Without one the block is not a table to any Markdown reader, and a model
    handed a pile of pipes reads it as prose.
    """
    if not rows:
        return []
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(padded[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def _section_furniture(document: Any) -> list[tuple[str, str]]:
    """Distinct header and footer text across the document's sections.

    Kept out of the body flow and appended at the end. Headers repeat on every
    page, so interleaving them would make a ten-page document read as though
    its title were restated between every paragraph — but dropping them loses
    document numbers, confidentiality markings, and revision codes, which in a
    contract or a policy are often the only provenance there is.
    """
    seen: set[tuple[str, str]] = set()
    out: list[tuple[str, str]] = []
    for section in document.sections:
        for label, source in (("header", section.header), ("footer", section.footer)):
            if source is None:
                continue
            try:
                text = " ".join(
                    t for t in (_paragraph_text(p) for p in source.paragraphs) if t
                ).strip()
            except (AttributeError, ValueError):
                continue
            if not text:
                continue
            key = (label, text)
            if key in seen:
                continue
            seen.add(key)
            out.append(key)
    return out


def extract_structured(raw: bytes) -> DocxExtract:
    """Read a ``.docx`` into Markdown, an outline, and its tables.

    Raises ``DocxReadError`` for bytes that are not a Word document. Callers
    inside best-effort pipelines should catch it; callers acting on a user's
    explicit upload should surface it, because "your file could not be read"
    is actionable and a silently empty extraction is not.
    """
    _require_python_docx()

    try:
        document = _new_document(io.BytesIO(raw))
    except Exception as exc:  # python-docx raises a variety of package errors
        raise DocxReadError(f"Not a readable Word document: {exc}") from exc

    numbering = _numbering_formats(document)

    md_lines: list[str] = []
    text_lines: list[str] = []
    outline: list[DocxHeading] = []
    tables: list[DocxTable] = []

    # Ordered lists count per (numId, level) so "1. 2. 3." survives instead of
    # becoming three items all numbered 1 — which reads to a model as three
    # copies of step one rather than a sequence.
    counters: dict[tuple[int | None, int], int] = {}
    # Consecutive monospace paragraphs are one fenced block, not one fence per
    # line, so indentation and blank lines inside a snippet survive.
    code_run: list[str] = []

    def blank() -> None:
        if md_lines and md_lines[-1] != "":
            md_lines.append("")

    def flush_code() -> None:
        if not code_run:
            return
        blank()
        md_lines.append("```")
        md_lines.extend(code_run)
        md_lines.append("```")
        md_lines.append("")
        text_lines.extend(code_run)
        code_run.clear()

    for block in _iter_block_items(document.element.body, document):
        if isinstance(block, _Table):
            flush_code()
            counters.clear()
            markdown_rows, plain_rows = _table_rows(block)
            if not markdown_rows:
                continue
            tables.append(DocxTable(rows=plain_rows))
            blank()
            md_lines.extend(_table_markdown(markdown_rows))
            md_lines.append("")
            text_lines.extend("\t".join(row) for row in plain_rows)
            continue

        if _is_code_paragraph(block):
            counters.clear()
            code_run.append(_paragraph_text(block))
            continue

        flush_code()

        plain = _paragraph_text(block)
        if not plain:
            if _is_horizontal_rule(block):
                counters.clear()
                blank()
                md_lines.append("---")
                md_lines.append("")
            continue

        markdown = _paragraph_markdown(block)

        level = _heading_level(block)
        if level is not None:
            counters.clear()
            blank()
            md_lines.append(f"{'#' * min(level, 6)} {markdown}")
            md_lines.append("")
            outline.append(DocxHeading(level=level, text=plain))
            text_lines.append(plain)
            continue

        marker = _list_marker(block, numbering)
        if marker is not None:
            indent_level, ordered, num_id = marker
            indent = "  " * min(indent_level, 5)
            if ordered:
                key = (num_id, indent_level)
                counters[key] = counters.get(key, 0) + 1
                # A nested level restarts each time its parent advances.
                for deeper in [k for k in counters if k[1] > indent_level]:
                    del counters[deeper]
                bullet = f"{counters[key]}."
            else:
                bullet = "-"
            md_lines.append(f"{indent}{bullet} {markdown}")
            text_lines.append(plain)
            continue

        counters.clear()

        if _is_quote(block):
            blank()
            md_lines.append(f"> {markdown}")
            md_lines.append("")
            text_lines.append(plain)
            continue

        blank()
        md_lines.append(markdown)
        md_lines.append("")
        text_lines.append(plain)

    flush_code()

    for label, text in _section_furniture(document):
        blank()
        md_lines.append(f"*Document {label}: {text}*")
        md_lines.append("")
        text_lines.append(text)

    markdown_doc = re.sub(r"\n{3,}", "\n\n", "\n".join(md_lines)).strip()
    plain_doc = "\n".join(text_lines).strip()

    return DocxExtract(
        markdown=markdown_doc,
        plain_text=plain_doc,
        outline=outline,
        tables=tables,
        paragraphs=_addressable_paragraphs(document),
        word_count=len(plain_doc.split()),
    )


def _addressable_paragraphs(document: Any) -> list[DocxParagraph]:
    """Every non-empty paragraph, in the order and the text the ops see.

    A second walk rather than a branch inside ``extract_structured``: that loop
    skips empties, folds monospace runs together and reads tables as
    rectangles, none of which is what an op addressing needs. Built from
    ``_all_paragraphs`` and ``_paragraph_text`` — the same two helpers
    ``PythonDocxAutomation.replace_text`` uses — because the point of this list
    is that a ``find`` checked against it is a ``find`` that will match.
    """
    out: list[DocxParagraph] = []
    for block in _iter_block_items(document.element.body, document):
        if isinstance(block, _Table):
            for row in block.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = _paragraph_text(paragraph)
                        if not text:
                            continue
                        out.append(
                            DocxParagraph(
                                index=len(out),
                                text=text,
                                style=_style_name(paragraph),
                                in_table=True,
                            )
                        )
            continue
        text = _paragraph_text(block)
        if not text:
            continue
        out.append(
            DocxParagraph(
                index=len(out),
                text=text,
                style=_style_name(block),
                heading_level=_heading_level(block),
            )
        )
    return out


# ─── Comments ─────────────────────────────────────────────────────────────
#
# Read straight off the zip rather than through python-docx, which models no
# part of this. Three parts have to agree before a comment means anything:
#
#   word/comments.xml          the remark, its author and its body
#   word/commentsExtended.xml  whether the thread is resolved, and who replies
#                              to whom (Word 2013+; older files omit it)
#   word/document.xml          the words the remark is ABOUT, delimited by
#                              w:commentRangeStart / w:commentRangeEnd
#
# The third is the one worth the effort. A comment carries no copy of the text
# it points at, so without walking the ranges a model reads "fix this" and has
# nothing to fix.

_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
_COMMENTS_PART = "word/comments.xml"
_COMMENTS_EX_PART = "word/commentsExtended.xml"
_DOCUMENT_PART = "word/document.xml"


def _w15(tag: str) -> str:
    """Clark notation for a ``w15:`` tag. python-docx's nsmap has no w15."""
    return f"{{{_W15}}}{tag}"


@dataclass(frozen=True)
class DocxComment:
    """One remark in a Word document, with the text it is attached to.

    ``id`` is the OOXML ``w:id``, which is unique within THIS package and
    nothing more: Word reuses an id after the comment holding it is deleted, so
    it is safe to address an op with and unsafe to remember across saves.

    ``resolved`` is false for a file that carries no ``commentsExtended`` part,
    which is the honest reading — an older Word could not mark a thread done.
    """

    id: str
    author: str
    initials: str
    text: str
    anchor_text: str
    date: str | None = None
    resolved: bool = False
    parent_id: str | None = None

    @property
    def is_reply(self) -> bool:
        return self.parent_id is not None


def _comment_anchors(document_xml: bytes) -> dict[str, str]:
    """The words each comment spans, by comment id.

    One walk in document order, carrying the set of ranges currently open, so a
    comment overlapping another still collects only its own text. A comment
    anchored at a point rather than over a range has no start/end pair at all
    and is simply absent here, which reads as an empty anchor.
    """
    try:
        root = parse_xml(document_xml)
    except Exception:  # noqa: BLE001 - malformed XML in any form means no anchors
        return {}

    start_tag = qn("w:commentRangeStart")
    end_tag = qn("w:commentRangeEnd")
    text_tag = qn("w:t")
    tab_tag = qn("w:tab")
    id_attr = qn("w:id")

    open_ids: set[str] = set()
    collected: dict[str, list[str]] = {}

    for node in root.iter():
        if node.tag == start_tag:
            comment_id = node.get(id_attr)
            if comment_id is not None:
                open_ids.add(comment_id)
                collected.setdefault(comment_id, [])
        elif node.tag == end_tag:
            open_ids.discard(node.get(id_attr) or "")
        elif open_ids and node.tag in (text_tag, tab_tag):
            piece = "\t" if node.tag == tab_tag else (node.text or "")
            for comment_id in open_ids:
                collected[comment_id].append(piece)

    return {key: "".join(pieces).strip() for key, pieces in collected.items()}


def extract_comments(raw: bytes) -> list[DocxComment]:
    """Every comment in a Word document, replies included, in document order.

    Returns an empty list for a document with no comments part — an absence,
    not a failure. Raises ``DocxReadError`` only when the bytes are not a zip
    at all, so a caller can tell "no comments" from "not a document".
    """
    _require_python_docx()

    try:
        archive = zipfile.ZipFile(io.BytesIO(raw))
    except Exception as exc:
        raise DocxReadError(f"Not a readable Word document: {exc}") from exc

    with archive:
        names = set(archive.namelist())
        if _COMMENTS_PART not in names:
            return []
        try:
            comments_xml = archive.read(_COMMENTS_PART)
            document_xml = (
                archive.read(_DOCUMENT_PART) if _DOCUMENT_PART in names else b""
            )
            extended_xml = (
                archive.read(_COMMENTS_EX_PART) if _COMMENTS_EX_PART in names else b""
            )
        except Exception as exc:
            raise DocxReadError(f"Word document is corrupt: {exc}") from exc

    try:
        root = parse_xml(comments_xml)
    except Exception as exc:
        raise DocxReadError(f"Word comments are unreadable: {exc}") from exc

    anchors = _comment_anchors(document_xml) if document_xml else {}

    # commentsExtended keys on the paraId of a comment's FIRST paragraph, not on
    # the comment id, so the two parts are joined through that.
    done_by_para: dict[str, bool] = {}
    parent_by_para: dict[str, str] = {}
    if extended_xml:
        try:
            for entry in parse_xml(extended_xml).iter(_w15("commentEx")):
                para_id = entry.get(_w15("paraId"))
                if not para_id:
                    continue
                done_by_para[para_id] = entry.get(_w15("done")) in ("1", "true")
                parent = entry.get(_w15("paraIdParent"))
                if parent:
                    parent_by_para[para_id] = parent
        except Exception:  # noqa: BLE001 - any parse failure has the same fallback
            # A malformed extended part costs threading and resolved state, not
            # the comments themselves.
            logger.warning("docx.commentsExtended_unreadable")

    text_tag = qn("w:t")
    paragraph_tag = qn("w:p")
    para_id_attr = qn("w14:paraId")

    para_id_of: dict[str, str] = {}
    comments: list[DocxComment] = []

    for element in root.iter(qn("w:comment")):
        comment_id = element.get(qn("w:id"))
        if comment_id is None:
            continue

        lines: list[str] = []
        for paragraph in element.iter(paragraph_tag):
            pieces = [node.text or "" for node in paragraph.iter(text_tag)]
            line = "".join(pieces).strip()
            if line:
                lines.append(line)

        first = next(element.iter(paragraph_tag), None)
        para_id = first.get(para_id_attr) if first is not None else None
        if para_id:
            para_id_of[comment_id] = para_id

        comments.append(
            DocxComment(
                id=comment_id,
                author=(element.get(qn("w:author")) or "").strip(),
                initials=(element.get(qn("w:initials")) or "").strip(),
                text="\n".join(lines),
                anchor_text=anchors.get(comment_id, ""),
                date=element.get(qn("w:date")),
                resolved=done_by_para.get(para_id or "", False),
            )
        )

    # Threading resolves in a second pass: a reply can appear before its parent
    # in the part, and paraId -> comment id is only complete once every comment
    # has been read.
    comment_id_by_para = {para: cid for cid, para in para_id_of.items()}
    threaded: list[DocxComment] = []
    for comment in comments:
        parent_para = parent_by_para.get(para_id_of.get(comment.id, ""))
        parent_id = comment_id_by_para.get(parent_para or "")
        if parent_id and parent_id != comment.id:
            comment = _dc_replace(comment, parent_id=parent_id)
        threaded.append(comment)
    return threaded


# ─── Rendering ────────────────────────────────────────────────────────────
#
# Renders the editor's own TipTap node vocabulary, which is also what
# `markdown_to_tiptap` produces. Tables are handled here but not there: they
# reach this renderer from real documents written in the editor, which has the
# table extension, rather than from Markdown.

_LINK_BLUE = "0563C1"
_MONOSPACE = "Consolas"

_TEMPLATES: dict[str, bytes] = {}

_TABLE_ROW = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_DIVIDER = re.compile(r"^\s*\|(?:\s*:?-{3,}:?\s*\|)+\s*$")


def _split_cells(line: str) -> list[str]:
    """Cells of a pipe-table row, honouring ``\\|`` as a literal pipe."""
    body = line.strip().strip("|")
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for char in body:
        if escaped:
            current.append(char if char == "|" else f"\\{char}")
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(char)
    if escaped:
        current.append("\\")
    cells.append("".join(current).strip())
    return cells


def _cell_paragraph(text: str) -> dict[str, Any]:
    """A table cell's text as a TipTap paragraph, inline marks included."""
    if not text.strip():
        return {"type": "paragraph"}
    try:
        parsed = markdown_to_tiptap(text)
    except MarkdownError:
        return {"type": "paragraph"}
    for node in parsed.get("content") or []:
        if node.get("type") == "paragraph":
            paragraph: dict[str, Any] = node
            return paragraph
    return {"type": "paragraph", "content": [{"type": "text", "text": text}]}


def _table_node(rows: list[list[str]]) -> dict[str, Any]:
    width = max(len(row) for row in rows)
    out_rows: list[dict[str, Any]] = []
    for index, row in enumerate(rows):
        kind = "tableHeader" if index == 0 else "tableCell"
        padded = row + [""] * (width - len(row))
        out_rows.append(
            {
                "type": "tableRow",
                "content": [
                    {"type": kind, "content": [_cell_paragraph(cell)]} for cell in padded
                ],
            }
        )
    return {"type": "table", "content": out_rows}


def _markdown_to_tree(markdown: str) -> dict[str, Any]:
    """Markdown to a TipTap document, pipe tables included.

    ``markdown_to_tiptap`` is the editor's write contract and deliberately has
    no tables, so it folds a pipe table into one run-together paragraph. That
    is fine for its own callers and wrong here for two reasons: extraction
    emits pipe tables, so a document could not survive its own round trip; and
    ``append_section`` takes Markdown written by a model, which reaches for a
    table the moment it has more than two columns of anything.

    So tables are carved out and built here, and every other construct is left
    to the one parser. Adding tables to that parser instead would change what
    the editor accepts over MCP, which is a separate decision.
    """
    lines = markdown.replace("\r\n", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    prose: list[str] = []
    index = 0

    def flush_prose() -> None:
        if not prose:
            return
        try:
            parsed = markdown_to_tiptap("\n".join(prose))
        except MarkdownError:
            prose.clear()
            return
        blocks.extend(parsed.get("content") or [])
        prose.clear()

    while index < len(lines):
        is_table = (
            _TABLE_ROW.match(lines[index])
            and index + 1 < len(lines)
            and _TABLE_DIVIDER.match(lines[index + 1])
        )
        if not is_table:
            prose.append(lines[index])
            index += 1
            continue

        flush_prose()
        rows = [_split_cells(lines[index])]
        index += 2  # header and divider
        while index < len(lines) and _TABLE_ROW.match(lines[index]):
            rows.append(_split_cells(lines[index]))
            index += 1
        blocks.append(_table_node(rows))

    flush_prose()

    if not blocks:
        raise DocxRenderError("That content produced an empty document.")

    return {"type": "doc", "content": blocks}


def _resolve_style(document: Any, *candidates: str) -> str | None:
    """The first of ``candidates`` this document's template actually defines.

    A document built from a corporate template may not have "List Bullet" or
    "Quote". Assigning a missing style raises, so an export of a perfectly
    ordinary document would fail on a styling detail nobody asked about.
    """
    try:
        available = {style.name for style in document.styles}
    except (AttributeError, ValueError):
        return None
    for name in candidates:
        if name in available:
            return name
    return None


def _add_hyperlink(paragraph: Any, url: str, text: str, *, bold: bool, italic: bool) -> None:
    """Append a real external hyperlink, not blue text that looks like one."""
    rel_id = paragraph.part.relate_to(url, _REL.HYPERLINK, is_external=True)

    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), _LINK_BLUE)
    rPr.append(colour)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    run.append(rPr)

    node = OxmlElement("w:t")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)

    link.append(run)
    paragraph._p.append(link)


def _write_inline(paragraph: Any, nodes: list[dict[str, Any]] | None) -> None:
    """Write TipTap inline content into a paragraph, carrying its marks."""
    for node in nodes or []:
        if node.get("type") != "text":
            continue
        text = node.get("text") or ""
        if not text:
            continue

        marks = {m.get("type") for m in node.get("marks") or []}
        href = next(
            (
                (m.get("attrs") or {}).get("href")
                for m in node.get("marks") or []
                if m.get("type") == "link"
            ),
            None,
        )

        if href:
            _add_hyperlink(
                paragraph, href, text, bold="bold" in marks, italic="italic" in marks
            )
            continue

        run = paragraph.add_run(text)
        run.bold = "bold" in marks or "strong" in marks
        run.italic = "italic" in marks or "em" in marks
        if "underline" in marks:
            run.underline = True
        if "code" in marks:
            run.font.name = _MONOSPACE


def _add_horizontal_rule(document: Any) -> None:
    paragraph = document.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    pPr.append(borders)


def _add_code_block(document: Any, node: dict[str, Any]) -> None:
    code = "".join(child.get("text") or "" for child in node.get("content") or [])
    style = _resolve_style(document, "No Spacing")
    for line in code.split("\n"):
        paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = _MONOSPACE
        run.font.size = Pt(9)


def _add_list(document: Any, node: dict[str, Any], ordered: bool, depth: int = 0) -> None:
    base = "List Number" if ordered else "List Bullet"
    # Word's built-in list styles are "List Bullet", "List Bullet 2", "List
    # Bullet 3" — indent is a different style, not a property.
    wanted = base if depth == 0 else f"{base} {min(depth + 1, 3)}"
    style = _resolve_style(document, wanted, base)

    for item in node.get("content") or []:
        if item.get("type") != "listItem":
            continue
        first = True
        for child in item.get("content") or []:
            kind = child.get("type")
            if kind in {"bulletList", "orderedList"}:
                _add_list(document, child, kind == "orderedList", depth + 1)
                continue
            paragraph = (
                document.add_paragraph(style=style) if style else document.add_paragraph()
            )
            if not first and not style:
                paragraph.paragraph_format.left_indent = Pt(18 * (depth + 1))
            _write_inline(paragraph, child.get("content"))
            first = False


def _add_table(document: Any, node: dict[str, Any]) -> None:
    rows = [r for r in (node.get("content") or []) if r.get("type") == "tableRow"]
    if not rows:
        return
    width = max(len(r.get("content") or []) for r in rows)
    if width == 0:
        return

    table = document.add_table(rows=0, cols=width)
    style = _resolve_style(document, "Table Grid")
    if style:
        table.style = style

    for row_node in rows:
        cells = row_node.get("content") or []
        row = table.add_row()
        is_header = any(cell.get("type") == "tableHeader" for cell in cells)
        if is_header:
            _mark_header_row(row)
        for index in range(width):
            cell = row.cells[index]
            # A fresh cell already holds one empty paragraph; write into it
            # rather than adding a second, which would double every row's
            # height.
            target = cell.paragraphs[0]
            if index >= len(cells):
                continue
            blocks = cells[index].get("content") or []
            for position, block in enumerate(blocks):
                paragraph = target if position == 0 else cell.add_paragraph()
                _write_inline(paragraph, block.get("content"))


def _mark_header_row(row: Any) -> None:
    """Mark a table row as a header row the way OOXML means it.

    ``w:tblHeader`` is the actual representation: Word repeats the row when the
    table breaks across pages. Bolding the runs instead would say the same
    thing to a reader while being a lie to Word — and it would make extraction
    non-idempotent, since re-extracting the rendered table would report
    ``**Tier**`` where the source said ``Tier``, so a document's stored text
    would drift every time it was saved.
    """
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


def _write_block(document: Any, node: dict[str, Any]) -> None:
    kind = node.get("type")

    if kind == "heading":
        level = int((node.get("attrs") or {}).get("level") or 1)
        style = _resolve_style(document, f"Heading {min(max(level, 1), 9)}")
        paragraph = (
            document.add_paragraph(style=style) if style else document.add_paragraph()
        )
        _write_inline(paragraph, node.get("content"))
    elif kind == "paragraph":
        _write_inline(document.add_paragraph(), node.get("content"))
    elif kind in {"bulletList", "orderedList"}:
        _add_list(document, node, kind == "orderedList")
    elif kind == "codeBlock":
        _add_code_block(document, node)
    elif kind == "blockquote":
        style = _resolve_style(document, "Quote", "Intense Quote")
        for child in node.get("content") or []:
            paragraph = (
                document.add_paragraph(style=style) if style else document.add_paragraph()
            )
            if not style:
                paragraph.paragraph_format.left_indent = Pt(36)
            _write_inline(paragraph, child.get("content"))
    elif kind == "horizontalRule":
        _add_horizontal_rule(document)
    elif kind == "table":
        _add_table(document, node)
    elif kind == "taskList":
        # The editor's checklists have no Word equivalent; a bulleted list
        # with an explicit box keeps the state visible instead of dropping it.
        for item in node.get("content") or []:
            checked = bool((item.get("attrs") or {}).get("checked"))
            style = _resolve_style(document, "List Bullet")
            for child in item.get("content") or []:
                paragraph = (
                    document.add_paragraph(style=style)
                    if style
                    else document.add_paragraph()
                )
                paragraph.add_run("☑ " if checked else "☐ ")
                _write_inline(paragraph, child.get("content"))
    else:
        # An unrecognised block still has text, and losing a section outright
        # is worse than rendering it as an unstyled paragraph.
        text = _collect_text(node)
        if text:
            document.add_paragraph(text)


def _collect_text(node: dict[str, Any]) -> str:
    if node.get("type") == "text":
        return node.get("text") or ""
    return "".join(_collect_text(child) for child in node.get("content") or [])


def render_docx(source: str | dict[str, Any], *, template_key: str | None = None) -> bytes:
    """Render Markdown or a TipTap document to ``.docx`` bytes.

    ``template_key`` selects a registered branded template to build on top of.
    An unknown key raises rather than silently producing an unbranded
    document, because "the export ignored our letterhead" is the kind of
    failure that is only noticed after it has been sent to a customer.
    """
    _require_python_docx()

    if isinstance(source, str):
        doc_tree = _markdown_to_tree(source)
    elif isinstance(source, dict):
        doc_tree = source
    else:
        raise DocxRenderError("Expected Markdown text or a TipTap document.")

    blocks = doc_tree.get("content") or []
    if not blocks:
        raise DocxRenderError("That content produced an empty document.")

    if template_key is None:
        document = _new_document()
    else:
        template = _TEMPLATES.get(template_key)
        if template is None:
            raise DocxRenderError(f"Unknown document template {template_key!r}.")
        document = _new_document(io.BytesIO(template))

    for node in blocks:
        _write_block(document, node)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ─── Automation ───────────────────────────────────────────────────────────


class DocxAutomationPort(Protocol):
    """Applying a structured edit list to a document's bytes.

    Async even though python-docx is synchronous. The callers are agent tools,
    Temporal activities, and the proposed-edit apply path — all async — and
    retrofitting ``async`` across those call sites later is the part that
    would actually hurt. A sidecar implementation is inherently async; making
    the port match it now costs a thread hop and keeps the swap to a config
    value.
    """

    async def apply_ops(
        self,
        raw: bytes,
        ops: list[dict[str, Any]],
        *,
        track_changes: bool = False,
    ) -> bytes:
        ...


# What may be PROPOSED, and what can be APPLIED here, are two different sets.
#
# They were one set, doing both jobs, and adding the comment ops to it would
# have silently let the headless backend accept ops python-docx cannot write —
# a proposal that validated cleanly and then half-applied.
PROPOSABLE_OPS = frozenset(
    {
        "replace_text",
        "set_table_cell",
        "append_section",
        "replace_section_body",
        # A remark rather than a rewrite: what a model should produce when it
        # has a concern it should not silently guess at. Reviewed as a comment
        # in the margin, not as a redline.
        "add_comment",
        "reply_to_comment",
        "resolve_comment",
    }
)

# What `PythonDocxAutomation` can write, which is the four text ops.
#
# The limit is python-docx, NOT the automation protocol: the Node automation
# host writes comments correctly, verified against saved bytes —
# `insertComment` produces `word/comments.xml` with the right author and a
# `w:commentRangeStart`/`End` pair, `replyToComment` threads through
# `commentsExtended`'s `w15:paraIdParent`, and `setCommentResolved` marks the
# whole thread `w15:done`. Only tracked changes are genuinely browser-only:
# that same host applies a plain edit with zero `w:ins`/`w:del` even against a
# package carrying `w:trackRevisions`. So adopting the sidecar moves the
# comment ops onto the unattended path too, and this set is what would shrink.
_HEADLESS_OPS = frozenset(
    {"replace_text", "set_table_cell", "append_section", "replace_section_body"}
)

# What each op needs to mean anything. Checked at propose time as well as at
# apply time: a proposal that cannot be applied should be refused when it is
# written, while the agent is still in a position to write a better one, not
# hours later when a reviewer clicks Apply.
_OP_REQUIRED_FIELDS: dict[str, tuple[str, ...]] = {
    "replace_text": ("find",),
    "set_table_cell": ("table_index", "row", "column"),
    "append_section": ("heading",),
    "replace_section_body": ("heading",),
    "add_comment": ("anchor_find", "text"),
    "reply_to_comment": ("comment_id", "text"),
    "resolve_comment": ("comment_id",),
}

# Ops the browser applies against a live editor and this backend refuses.
BROWSER_ONLY_OPS = PROPOSABLE_OPS - _HEADLESS_OPS

# Ops carrying a coordinate the browser cannot work out for itself, and so need
# stamping against the bytes before a reviewer sees them.
_COMMENT_TARGET_OPS = frozenset({"reply_to_comment", "resolve_comment"})
_RESOLVABLE_OPS = frozenset({"set_table_cell", "add_comment"}) | _COMMENT_TARGET_OPS


def _comment_label(comment: DocxComment) -> str:
    """What to call a comment in the review UI.

    Its own text, trimmed, because that is how the person who wrote it will
    recognise it — "comment 3" names nothing a reader can see.
    """
    body = " ".join(comment.text.split())
    if len(body) > 60:
        body = body[:57].rstrip() + "…"
    author = comment.author or "someone"
    return f"{author}: “{body}”" if body else f"a comment by {author}"


def resolve_ops_for_review(
    raw: bytes, ops: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    """Stamp each op with whatever the browser cannot work out for itself.

    Today that is one thing, and it is the difference between a table edit being
    reviewable and not. The editor's automation protocol has **no table
    operations** — no enumeration, no cell addressing — so a
    ``(table_index, row, column)`` coordinate is unresolvable in the browser. It
    is trivially resolvable here, because this side holds the bytes.

    So a ``set_table_cell`` op gains ``expected_current`` (what the cell says now)
    and ``cell_label`` (what to call it in the UI). The browser locates the cell
    by searching for that text, which also makes the edit refuse itself if the
    cell has changed since — a staleness check at cell granularity.

    A cell whose text the browser could not locate is left **unresolved with a
    stated reason** rather than resolved to something that will not match. Two
    cases: an out-of-range coordinate, and a cell holding more than one
    paragraph — the latter because its text spans a paragraph break, which a
    document-text search cannot match, and "no longer says X" would be a
    misleading way to report a shape that was never supported.

    The comment ops get the same treatment for the same reason. A ``w:id`` is
    unique within one package and nothing more — Word reuses an id once the
    comment holding it is deleted — so an op saying "reply to comment 3" is a
    coordinate as fragile as a table cell's. Stamping the author and text the
    comment had when the AI read it lets the browser refuse to reply to a
    different remark that happens to have inherited the number.

    Returns a new list; the stored ops are never mutated in place, because the
    resolution is only valid against the bytes it was computed from.
    """
    if not any(op.get("kind") in _RESOLVABLE_OPS for op in ops):
        return [dict(op) for op in ops]

    try:
        document = _new_document(io.BytesIO(raw))
        tables = _body_tables(document)
        paragraph_text = [p.text for p in _addressable_paragraphs(document)]
    except Exception as exc:  # noqa: BLE001 - any package-level failure
        logger.warning("Could not read tables while resolving ops: %s", exc)
        tables = []
        paragraph_text = []

    comments_by_id: dict[str, DocxComment] = {}
    if any(op.get("kind") in _COMMENT_TARGET_OPS for op in ops):
        try:
            comments_by_id = {c.id: c for c in extract_comments(raw)}
        except DocxReadError as exc:
            logger.warning("Could not read comments while resolving ops: %s", exc)

    resolved: list[dict[str, Any]] = []
    for op in ops:
        kind = op.get("kind")

        if kind == "add_comment":
            enriched = dict(op)
            anchor = op.get("anchor_find")
            if not isinstance(anchor, str) or not anchor:
                enriched["unresolvable"] = "that comment has nothing to attach to"
            elif paragraph_text and not any(anchor in text for text in paragraph_text):
                enriched["unresolvable"] = (
                    f"the text to comment on ({anchor!r}) is not in this document"
                )
            resolved.append(enriched)
            continue

        if kind in _COMMENT_TARGET_OPS:
            enriched = dict(op)
            target = comments_by_id.get(str(op.get("comment_id")))
            if target is None:
                enriched["unresolvable"] = (
                    "that comment is no longer in this document"
                )
            else:
                enriched["expected_comment_author"] = target.author
                enriched["expected_comment_text"] = target.text
                enriched["cell_label"] = _comment_label(target)
            resolved.append(enriched)
            continue

        if kind != "set_table_cell":
            resolved.append(dict(op))
            continue

        enriched = dict(op)
        try:
            index = int(op["table_index"])
            row_index = int(op["row"])
            column = int(op["column"])
            cell = tables[index].rows[row_index].cells[column]
        except (KeyError, TypeError, ValueError, IndexError):
            enriched["unresolvable"] = "that cell coordinate is not in this document"
            resolved.append(enriched)
            continue

        paragraphs = [p for p in cell.paragraphs if _paragraph_text(p)]
        if len(paragraphs) > 1:
            enriched["unresolvable"] = (
                "that cell holds more than one paragraph, which cannot be "
                "located by its text"
            )
            resolved.append(enriched)
            continue

        enriched["expected_current"] = _paragraph_text(paragraphs[0]) if paragraphs else ""
        header = [_cell_plain(c) for c in tables[index].rows[0].cells] if tables[index].rows else []
        column_name = (
            header[column]
            if column < len(header) and header[column]
            else f"column {column + 1}"
        )
        enriched["cell_label"] = f"{column_name}, row {row_index + 1}"
        resolved.append(enriched)

    return resolved


def validate_ops(ops: list[dict[str, Any]]) -> None:
    """Check an op list's shape without needing the document.

    Raises ``DocxOpUnsupported`` naming the offending index, because an agent
    handed "invalid ops" learns nothing, and a list of ten with one bad entry is
    exactly where a precise message saves a round trip.
    """
    if not ops:
        raise DocxOpUnsupported("An edit needs at least one op.")

    for index, op in enumerate(ops):
        if not isinstance(op, dict):
            raise DocxOpUnsupported(f"Op {index} is not an object.")
        kind = op.get("kind")
        if kind not in PROPOSABLE_OPS:
            raise DocxOpUnsupported(
                f"Op {index} has kind {kind!r}; supported kinds are "
                f"{sorted(PROPOSABLE_OPS)}."
            )
        for field_name in _OP_REQUIRED_FIELDS[kind]:
            if op.get(field_name) is None:
                raise DocxOpUnsupported(
                    f"Op {index} ({kind}) is missing required field {field_name!r}."
                )


def _require(op: dict[str, Any], key: str) -> Any:
    if key not in op or op[key] is None:
        raise DocxOpUnsupported(f"Op {op.get('kind')!r} is missing required field {key!r}.")
    return op[key]


def _replace_in_paragraph(paragraph: Any, find: str, replace: str) -> int:
    """Replace every occurrence of ``find`` within one paragraph.

    A single word is routinely split across several ``w:r`` runs — Word starts
    a new run at a spellcheck boundary, a formatting change, or a saved
    revision — so matching run by run misses most real occurrences. This joins
    the runs, finds matches in the joined text, and writes the result back.

    When a match spans runs, the replacement takes the formatting of the run
    the match started in and the rest of the span is emptied. Stated because
    it is a real behaviour: replacing a phrase whose second half was bold
    yields an unbold replacement.
    """
    runs = [
        node for node in paragraph._p.iterchildren() if node.tag == qn("w:r")
    ]
    if not runs:
        return 0

    texts: list[str] = []
    for run in runs:
        nodes = run.findall(qn("w:t"))
        texts.append("".join(node.text or "" for node in nodes))

    joined = "".join(texts)
    if find not in joined:
        return 0

    # Offsets of each run within the joined text, so the rewrite can be
    # attributed back to the run each character came from. Distributing by
    # original run *length* instead would corrupt every boundary whenever the
    # replacement is a different length than the match.
    spans: list[tuple[int, int]] = []
    cursor = 0
    for text in texts:
        spans.append((cursor, cursor + len(text)))
        cursor += len(text)

    rewritten = ["" for _ in runs]
    count = 0
    position = 0
    index = 0

    while position < len(joined):
        # Advance to the run owning this character, stepping over empty runs.
        while index < len(spans) - 1 and position >= spans[index][1]:
            index += 1

        if joined.startswith(find, position):
            # The whole replacement lands in the run the match started in;
            # characters the match consumed from later runs simply vanish.
            rewritten[index] += replace
            position += len(find)
            count += 1
        else:
            rewritten[index] += joined[position]
            position += 1

    for run, text in zip(runs, rewritten, strict=True):
        _set_run_text(run, text)

    return count


def _set_run_text(run: Any, text: str) -> None:
    """Make a run hold exactly ``text``, dropping its extra ``w:t`` nodes."""
    nodes = run.findall(qn("w:t"))
    if not nodes:
        if not text:
            return
        node = OxmlElement("w:t")
        run.append(node)
        nodes = [node]
    nodes[0].text = text
    nodes[0].set(qn("xml:space"), "preserve")
    for extra in nodes[1:]:
        run.remove(extra)


def _all_paragraphs(document: Any) -> list[Any]:
    """Every paragraph in the document, table cells included."""
    out: list[Any] = []
    for block in _iter_block_items(document.element.body, document):
        if isinstance(block, _Table):
            for row in block.rows:
                for cell in row.cells:
                    out.extend(cell.paragraphs)
        else:
            out.append(block)
    return out


def _body_tables(document: Any) -> list[Any]:
    return [
        block
        for block in _iter_block_items(document.element.body, document)
        if isinstance(block, _Table)
    ]


def _rendered_blocks(markdown: str) -> list[Any]:
    """Body children of a document rendered from ``markdown``.

    Rendering to bytes and reading them back is deliberate: it means inserted
    content goes through exactly the same path as ``render_docx``, so a
    section appended by an agent is styled identically to one exported by a
    user.
    """
    rendered = render_docx(markdown)
    scratch = _new_document(io.BytesIO(rendered))
    return [copy.deepcopy(child) for child in scratch.element.body.iterchildren()]


def _is_section_end(element: Any, document: Any, level: int) -> bool:
    """Whether ``element`` starts a heading at or above ``level``."""
    if element.tag != qn("w:p"):
        return False
    found = _heading_level(_Paragraph(element, document))
    return found is not None and found <= level


def _find_heading(document: Any, title: str) -> tuple[Any, int] | None:
    """The first heading paragraph with this text, and its level.

    Returned as a pair because the paragraph is useless without its level:
    the level is what decides where the section it opens ends.
    """
    for block in _iter_block_items(document.element.body, document):
        if isinstance(block, _Table):
            continue
        level = _heading_level(block)
        if level is not None and _paragraph_text(block).strip() == title:
            return block, level
    return None


class PythonDocxAutomation:
    """The restricted, in-process automation backend.

    Honest about its limits: it refuses tracked changes outright, and refuses
    any op outside ``_HEADLESS_OPS`` — the comment ops included — rather than
    approximating it. Both refusals are ``DocxOpUnsupported``, which is the
    signal to watch; see that exception's docstring.
    """

    supports_tracked_changes = False

    async def apply_ops(
        self,
        raw: bytes,
        ops: list[dict[str, Any]],
        *,
        track_changes: bool = False,
    ) -> bytes:
        if track_changes:
            self._refuse(
                "tracked changes",
                "python-docx cannot write w:ins/w:del; use the sidecar backend "
                "or apply this edit through the browser editor",
            )

        _require_python_docx()

        try:
            document = _new_document(io.BytesIO(raw))
        except Exception as exc:
            raise DocxReadError(f"Not a readable Word document: {exc}") from exc

        for op in ops:
            kind = op.get("kind")
            if kind in BROWSER_ONLY_OPS:
                self._refuse(
                    f"op {kind!r}",
                    "python-docx cannot write comments; this op needs the "
                    "browser editor or the sidecar backend",
                )
            if kind not in _HEADLESS_OPS:
                self._refuse(
                    f"op {kind!r}",
                    f"this backend supports {sorted(_HEADLESS_OPS)}",
                )
            getattr(self, f"_op_{kind}")(document, op)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _refuse(what: str, why: str) -> NoReturn:
        logger.warning("docx.op_unsupported backend=python-docx what=%s why=%s", what, why)
        raise DocxOpUnsupported(f"Cannot apply {what}: {why}.")

    # ── ops ──

    def _op_replace_text(self, document: Any, op: dict[str, Any]) -> None:
        find = _require(op, "find")
        replace = op.get("replace") or ""
        if not isinstance(find, str) or not find:
            raise DocxOpUnsupported("replace_text needs a non-empty 'find' string.")

        replaced = 0
        for paragraph in _all_paragraphs(document):
            replaced += _replace_in_paragraph(paragraph, find, replace)

        if replaced == 0:
            self._refuse(
                f"replace_text {find!r}",
                "the text does not appear in this document, so the edit would be a "
                "silent no-op",
            )

        expected = op.get("count")
        if expected is not None and replaced != expected:
            self._refuse(
                f"replace_text {find!r}",
                f"expected {expected} occurrence(s) but found {replaced}",
            )

    def _op_set_table_cell(self, document: Any, op: dict[str, Any]) -> None:
        tables = _body_tables(document)
        index = int(_require(op, "table_index"))
        row_index = int(_require(op, "row"))
        column = int(_require(op, "column"))
        text = op.get("text") or ""

        if not 0 <= index < len(tables):
            self._refuse(
                "set_table_cell",
                f"table_index {index} is out of range ({len(tables)} table(s))",
            )
        table = tables[index]
        if not 0 <= row_index < len(table.rows):
            self._refuse(
                "set_table_cell",
                f"row {row_index} is out of range ({len(table.rows)} row(s))",
            )
        row = table.rows[row_index]
        if not 0 <= column < len(row.cells):
            self._refuse(
                "set_table_cell",
                f"column {column} is out of range ({len(row.cells)} column(s))",
            )

        cell = row.cells[column]
        first = cell.paragraphs[0]
        for extra in cell.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
        for run in list(first._p.iterchildren()):
            if run.tag in (qn("w:r"), qn("w:hyperlink")):
                first._p.remove(run)
        first.add_run(text)

    def _op_append_section(self, document: Any, op: dict[str, Any]) -> None:
        heading = _require(op, "heading")
        level = int(op.get("level") or 2)
        body = op.get("markdown") or ""
        markdown = f"{'#' * min(max(level, 1), 6)} {heading}\n\n{body}".strip()
        for element in _rendered_blocks(markdown):
            document.element.body.append(element)

    def _op_replace_section_body(self, document: Any, op: dict[str, Any]) -> None:
        heading = _require(op, "heading")
        body = op.get("markdown") or ""

        found = _find_heading(document, str(heading).strip())
        if found is None:
            self._refuse(
                "replace_section_body",
                f"no heading titled {heading!r} in this document",
            )
        anchor, level = found

        # Drop everything between this heading and the next one at the same or
        # a higher level; the heading itself stays.
        cursor = anchor._p.getnext()
        while cursor is not None and not _is_section_end(cursor, document, level):
            doomed, cursor = cursor, cursor.getnext()
            doomed.getparent().remove(doomed)

        if body.strip():
            previous = anchor._p
            for element in _rendered_blocks(body):
                previous.addnext(element)
                previous = element


_automation: DocxAutomationPort | None = None


def get_docx_automation() -> DocxAutomationPort:
    """The configured automation backend.

    One implementation today. The indirection is the seam a Node sidecar
    arrives behind — see this module's docstring.
    """
    global _automation
    if _automation is None:
        _automation = PythonDocxAutomation()
    return _automation
